"""
Knowledge Management API Module

This module handles all knowledge base operations including:
- Crawling and indexing web content
- Document upload and processing  
- RAG (Retrieval Augmented Generation) queries
- Knowledge item management and search
- Real-time progress tracking via WebSockets
"""

import asyncio
import json
import os
import secrets
import uuid
from datetime import datetime
from typing import Any, Dict, List, Optional
from fastapi import APIRouter, HTTPException, WebSocket, WebSocketDisconnect, UploadFile, File, Form, Request
from pydantic import BaseModel, HttpUrl
from pathlib import Path
import traceback

from ..utils import get_supabase_client, add_documents_to_supabase, extract_source_summary
from ..services.rag.document_storage_service import DocumentStorageService

# Import Logfire - use logfire directly like other working APIs
from ..logfire_config import logfire

# Document processing imports
try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False
    
try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False
    
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

import io

# Create router
router = APIRouter(prefix="/api", tags=["knowledge"])

# Feature flag for Socket.IO migration
USE_SOCKETIO_PROGRESS = os.getenv("USE_SOCKETIO_PROGRESS", "true").lower() == "true"

# Get the global crawling context from main.py
def get_crawling_context():
    """Get the global crawling context from the app state."""
    from fastapi import Request
    import inspect
    
    # Try to get the context from the current request
    frame = inspect.currentframe()
    try:
        # Walk up the call stack to find the app instance
        while frame:
            if 'app' in frame.f_locals and hasattr(frame.f_locals['app'], 'state'):
                app = frame.f_locals['app']
                if hasattr(app.state, 'crawling_context'):
                    return app.state.crawling_context
            frame = frame.f_back
    finally:
        del frame
    
    # Fallback - import directly from main
    try:
        from ..main import crawling_context
        return crawling_context
    except ImportError:
        # Create a minimal context if we can't get the global one
        class MinimalContext:
            def __init__(self):
                self._initialized = False
                self.supabase_client = get_supabase_client()
        return MinimalContext()

# Connection Manager for WebSocket management
class ConnectionManager:
    def __init__(self):
        self.active_connections: List[WebSocket] = []
    
    async def connect(self, websocket: WebSocket):
        await websocket.accept()
        self.active_connections.append(websocket)
    
    def disconnect(self, websocket: WebSocket):
        if websocket in self.active_connections:
            self.active_connections.remove(websocket)
    
    async def send_personal_message(self, message: str, websocket: WebSocket):
        try:
            await websocket.send_text(message)
        except:
            self.disconnect(websocket)
    
    async def broadcast(self, message: dict):
        disconnected = []
        for connection in self.active_connections:
            try:
                await connection.send_json(message)
            except:
                disconnected.append(connection)
        
        # Remove disconnected clients
        for conn in disconnected:
            self.disconnect(conn)

# Global connection manager
manager = ConnectionManager()

# Request Models
class KnowledgeItemRequest(BaseModel):
    url: str
    knowledge_type: str = 'technical'
    tags: List[str] = []
    update_frequency: int = 7
    max_depth: int = 2  # Maximum crawl depth (1-5)

class CrawlRequest(BaseModel):
    url: str
    knowledge_type: str = 'general'
    tags: List[str] = []
    update_frequency: int = 7
    max_depth: int = 2  # Maximum crawl depth (1-5)

class RagQueryRequest(BaseModel):
    query: str
    source: Optional[str] = None
    match_count: int = 5

# Progress Manager for crawling operations
class CrawlProgressManager:
    """Manages crawling progress tracking and WebSocket streaming."""
    
    def __init__(self):
        self.active_crawls: Dict[str, Dict[str, Any]] = {}
        self.progress_websockets: Dict[str, List[WebSocket]] = {}
        self.connection_ready_events: Dict[str, asyncio.Event] = {}
    
    def start_crawl(self, progress_id: str, initial_data: Dict[str, Any]) -> None:
        """Start tracking a new crawl operation."""
        # Ensure logs is always a list
        if 'logs' not in initial_data:
            initial_data['logs'] = []
        elif isinstance(initial_data.get('logs'), str):
            initial_data['logs'] = [initial_data['logs']]
            
        self.active_crawls[progress_id] = {
            'status': 'starting',
            'percentage': 0,
            'start_time': datetime.now(),
            'logs': initial_data.get('logs', ['Starting crawl...']),
            **initial_data
        }
        
        # Create event for this progress ID
        self.connection_ready_events[progress_id] = asyncio.Event()
        
    async def update_progress(self, progress_id: str, update_data: Dict[str, Any]) -> None:
        """Update crawling progress and notify connected clients."""
        if progress_id not in self.active_crawls:
            return
        
        # Update progress data
        self.active_crawls[progress_id].update(update_data)
        self.active_crawls[progress_id]['updated_at'] = datetime.now()
        
        # Add to logs if message provided
        if 'log' in update_data:
            self.active_crawls[progress_id]['logs'].append(update_data['log'])
        
        # Broadcast to connected clients
        await self._broadcast_progress(progress_id)
    
    async def complete_crawl(self, progress_id: str, completion_data: Dict[str, Any]) -> None:
        """Mark crawl as completed and notify clients."""
        if progress_id not in self.active_crawls:
            return
        
        self.active_crawls[progress_id].update({
            'status': 'completed',
            'percentage': 100,
            'completed_at': datetime.now(),
            **completion_data
        })
        
        if 'log' in completion_data:
            self.active_crawls[progress_id]['logs'].append(completion_data['log'])
        
        await self._broadcast_progress(progress_id)
        
        # Clean up after 5 minutes
        await asyncio.sleep(300)
        if progress_id in self.active_crawls:
            del self.active_crawls[progress_id]
        if progress_id in self.connection_ready_events:
            del self.connection_ready_events[progress_id]
    
    async def error_crawl(self, progress_id: str, error_message: str) -> None:
        """Mark crawl as failed and notify clients."""
        if progress_id not in self.active_crawls:
            return
        
        self.active_crawls[progress_id].update({
            'status': 'error',
            'error': error_message,
            'completed_at': datetime.now()
        })
        
        self.active_crawls[progress_id]['logs'].append(f"Error: {error_message}")
        
        await self._broadcast_progress(progress_id)
    
    async def add_websocket(self, progress_id: str, websocket: WebSocket) -> None:
        """Add WebSocket to progress tracking."""
        try:
            print(f"DEBUG: add_websocket called for progress_id: {progress_id}")
            await websocket.accept()
            print(f"DEBUG: WebSocket accepted for progress_id: {progress_id}")
            
            if progress_id not in self.progress_websockets:
                self.progress_websockets[progress_id] = []
            
            self.progress_websockets[progress_id].append(websocket)
            print(f"DEBUG: WebSocket added to progress_websockets for {progress_id}. Total websockets: {len(self.progress_websockets[progress_id])}")
            
            # Signal that connection is ready
            if progress_id in self.connection_ready_events:
                self.connection_ready_events[progress_id].set()
                print(f"DEBUG: Connection ready event set for {progress_id}")
            
            # Send current progress immediately if available
            if progress_id in self.active_crawls:
                print(f"DEBUG: Found active crawl for {progress_id}, sending current state")
                try:
                    # Ensure progressId is included in the data
                    data = self.active_crawls[progress_id].copy()
                    data['progressId'] = progress_id
                    
                    # Convert ALL datetime objects to strings for JSON serialization
                    for key, value in data.items():
                        if hasattr(value, 'isoformat'):
                            data[key] = value.isoformat()
                    
                    message = {
                        "type": "crawl_progress",
                        "data": data
                    }
                    print(f"DEBUG: Sending initial progress to new WebSocket: {message}")
                    await websocket.send_json(message)
                except Exception as e:
                    print(f"DEBUG: Error sending initial progress: {e}")
            else:
                print(f"DEBUG: No active crawl found for {progress_id} in active_crawls. Keys: {list(self.active_crawls.keys())}")
                # Send a waiting message
                try:
                    await websocket.send_json({
                        "type": "crawl_progress",
                        "data": {
                            "progressId": progress_id,
                            "status": "waiting",
                            "percentage": 0,
                            "logs": ["Waiting for crawl to start..."]
                        }
                    })
                except Exception as e:
                    print(f"DEBUG: Error sending waiting message: {e}")
        except Exception as e:
            print(f"ERROR: Exception in add_websocket: {e}")
            import traceback
            traceback.print_exc()
            raise
    
    def remove_websocket(self, progress_id: str, websocket: WebSocket) -> None:
        """Remove WebSocket from progress tracking."""
        if progress_id in self.progress_websockets:
            try:
                self.progress_websockets[progress_id].remove(websocket)
                if not self.progress_websockets[progress_id]:
                    del self.progress_websockets[progress_id]
            except ValueError:
                pass
    
    async def wait_for_websocket_connection(self, progress_id: str, timeout: float = 5.0) -> bool:
        """Wait for WebSocket connection to be established."""
        if progress_id not in self.connection_ready_events:
            print(f"WARNING: No connection event found for {progress_id}")
            return False
        
        try:
            print(f"DEBUG: Waiting for WebSocket connection for {progress_id} (timeout: {timeout}s)")
            await asyncio.wait_for(self.connection_ready_events[progress_id].wait(), timeout=timeout)
            print(f"DEBUG: WebSocket connection established for {progress_id}")
            return True
        except asyncio.TimeoutError:
            print(f"WARNING: WebSocket connection timeout for {progress_id} after {timeout}s")
            return False
    
    async def _broadcast_progress(self, progress_id: str) -> None:
        """Broadcast progress update to all connected clients."""
        progress_data = self.active_crawls.get(progress_id, {}).copy()
        progress_data['progressId'] = progress_id
        progress_data['progress_id'] = progress_id  # Add both formats for compatibility
        
        # Convert ALL datetime objects to strings for JSON serialization
        for key, value in progress_data.items():
            if hasattr(value, 'isoformat'):
                progress_data[key] = value.isoformat()
        
        # Ensure all values are JSON serializable
        try:
            json.dumps(progress_data)  # Test serialization
        except (TypeError, ValueError) as e:
            print(f"DEBUG: JSON serialization error for progress_id {progress_id}: {e}")
            # Create a fallback safe message
            progress_data = {
                'progressId': progress_id,
                'status': str(progress_data.get('status', 'unknown')),
                'percentage': progress_data.get('percentage', 0),
                'logs': progress_data.get('logs', [])
            }
        
        message = {
            "type": "crawl_progress" if progress_data.get('status') != 'completed' else "crawl_completed",
            "data": progress_data
        }
        
        # Emit via Socket.IO
        try:
            event_type = 'progress_update' if progress_data.get('status') != 'completed' else 'progress_complete'
            await sio.emit(event_type, progress_data, room=progress_id, namespace=NAMESPACE_CRAWL)
            print(f"DEBUG: Emitted {event_type} via Socket.IO for {progress_id}")
        except Exception as e:
            print(f"DEBUG: Failed to emit via Socket.IO: {e}")
        
        # Keep legacy WebSocket support for now
        if progress_id in self.progress_websockets:
            print(f"DEBUG: Broadcasting to {len(self.progress_websockets[progress_id])} WebSocket(s) for {progress_id}: {progress_data.get('status')} {progress_data.get('percentage')}%")
            
            # Send to all connected WebSocket clients with improved error handling
            disconnected = []
            successful_sends = 0
            
            for websocket in self.progress_websockets[progress_id]:
                try:
                    await websocket.send_json(message)
                    successful_sends += 1
                    print(f"DEBUG: Successfully sent progress update to WebSocket")
                except Exception as e:
                    print(f"DEBUG: Failed to send to WebSocket: {e}")
                    disconnected.append(websocket)
            
            # Clean up disconnected WebSockets
            for ws in disconnected:
                self.remove_websocket(progress_id, ws)
            
            print(f"DEBUG: Broadcast completed: {successful_sends} successful, {len(disconnected)} failed")
            
            # If all WebSockets failed, log warning
            if successful_sends == 0 and len(self.progress_websockets.get(progress_id, [])) > 0:
                print(f"WARNING: All WebSocket connections failed for progress_id: {progress_id}")

# Global progress manager - now always uses Socket.IO under the hood
progress_manager = CrawlProgressManager()
if USE_SOCKETIO_PROGRESS:
    logfire.info("Using Socket.IO for real-time progress updates")
else:
    logfire.info("Using legacy WebSocket mode (deprecated)")

async def get_available_sources_direct() -> Dict[str, Any]:
    """Get all available sources from the sources table directly."""
    try:
        supabase_client = get_supabase_client()
        
        # Query the sources table directly
        result = supabase_client.from_('sources')\
            .select('*')\
            .order('source_id')\
            .execute()
        
        # Format the sources with their details
        sources = []
        if result.data:
            for source in result.data:
                sources.append({
                    "source_id": source.get("source_id"),
                    "title": source.get("title", source.get("summary", "Untitled")),
                    "summary": source.get("summary"),
                    "metadata": source.get("metadata", {}),
                    "total_words": source.get("total_words", source.get("total_word_count", 0)),
                    "update_frequency": source.get("update_frequency", 7),  # Include frequency from sources table
                    "created_at": source.get("created_at"),
                    "updated_at": source.get("updated_at", source.get("created_at"))
                })
        
        return {
            "success": True,
            "sources": sources,
            "count": len(sources)
        }
    except Exception as e:
        return {
            "success": False,
            "error": str(e)
        }

@router.get("/knowledge-items")
async def get_knowledge_items(
    page: int = 1,
    per_page: int = 20,
    knowledge_type: Optional[str] = None,
    search: Optional[str] = None
):
    """Get knowledge items with pagination and filtering."""
    with logfire.span("api_get_knowledge_items") as span:
        span.set_attribute("endpoint", "/api/knowledge-items")
        span.set_attribute("method", "GET")
        span.set_attribute("page", page)
        span.set_attribute("per_page", per_page)
        span.set_attribute("knowledge_type", knowledge_type)
        span.set_attribute("has_search", search is not None)
        
        try:
            logfire.info("Getting knowledge items", page=page, per_page=per_page, knowledge_type=knowledge_type, search=search)
            
            # Ensure crawling context is initialized once  
            crawling_context = get_crawling_context()
            if not crawling_context._initialized:
                await crawling_context.initialize()
            
            # Get all sources directly 
            sources_result = await get_available_sources_direct()
            
            # Parse the JSON response
            if isinstance(sources_result, str):
                sources_data = json.loads(sources_result)
            else:
                sources_data = sources_result
            
            # Transform the data to match frontend expectations
            items = []
            for source in sources_data.get('sources', []):
                # Use title and metadata from sources table
                source_metadata = source.get('metadata', {})
                source_id = source['source_id']
                
                # Get first page URL if available - simplified query
                supabase_client = crawling_context.supabase_client or get_supabase_client()
                try:
                    pages_response = supabase_client.from_('crawled_pages')\
                        .select('url')\
                        .eq('source_id', source_id)\
                        .limit(1)\
                        .execute()
                    
                    first_page = pages_response.data[0] if pages_response.data else {}
                    first_page_url = first_page.get('url', f"source://{source_id}")
                except:
                    first_page_url = f"source://{source_id}"
                
                # Determine source type - if metadata has source_type='file', use it; otherwise check URL pattern
                stored_source_type = source_metadata.get('source_type')
                if stored_source_type:
                    source_type = stored_source_type
                else:
                    # Legacy fallback - check URL pattern
                    source_type = 'file' if first_page_url.startswith('file://') else 'url'
                
                # Get code examples for this source - simplified query for now
                code_examples = []
                try:
                    code_examples_response = supabase_client.from_('code_examples')\
                        .select('id, summary, metadata')\
                        .eq('source_id', source_id)\
                        .execute()
                    
                    code_examples = code_examples_response.data if code_examples_response.data else []
                except:
                    code_examples = []
                
                item = {
                    'id': source_id,
                    'title': source.get('title', source.get('summary', 'Untitled')),
                    'url': first_page_url,
                    'source_id': source_id,
                    'code_examples': code_examples,  # Include code examples data
                    'metadata': {
                        'knowledge_type': source_metadata.get('knowledge_type', 'technical'),
                        'tags': source_metadata.get('tags', []),
                        'source_type': source_type,
                        'status': 'active',
                        'description': source_metadata.get('description', source.get('summary', '')),
                        'chunks_count': source.get('total_words', 0),
                        'word_count': source.get('total_words', 0),
                        'last_scraped': source.get('updated_at'),
                        'file_name': source_metadata.get('file_name'),
                        'file_type': source_metadata.get('file_type'),
                        'update_frequency': source.get('update_frequency', 7),  # Include frequency from sources table
                        'code_examples_count': len(code_examples),  # Add count for easy access
                        **source_metadata
                    },
                    'created_at': source.get('created_at'),
                    'updated_at': source.get('updated_at')
                }
                items.append(item)
            
            # Filter by search term if provided
            if search:
                search_lower = search.lower()
                items = [
                    item for item in items 
                    if search_lower in item['title'].lower() 
                    or search_lower in item['metadata'].get('description', '').lower()
                    or any(search_lower in tag.lower() for tag in item['metadata'].get('tags', []))
                ]
            
            # Filter by knowledge type if provided
            if knowledge_type:
                items = [
                    item for item in items 
                    if item['metadata'].get('knowledge_type') == knowledge_type
                ]
            
            # Apply pagination
            total = len(items)
            start_idx = (page - 1) * per_page
            end_idx = start_idx + per_page
            paginated_items = items[start_idx:end_idx]
            
            logfire.info("Knowledge items retrieved", total=total, page=page, filtered_count=len(paginated_items))
            span.set_attribute("total_items", total)
            span.set_attribute("filtered_items", len(paginated_items))
            
            return {
                'items': paginated_items,
                'total': total,
                'page': page,
                'per_page': per_page,
                'pages': (total + per_page - 1) // per_page
            }
            
        except Exception as e:
            logfire.error("Failed to get knowledge items", error=str(e), page=page, per_page=per_page)
            span.set_attribute("error", str(e))
            raise HTTPException(status_code=500, detail={'error': str(e)})

@router.put("/knowledge-items/{source_id}")
async def update_knowledge_item(source_id: str, updates: dict):
    """Update a knowledge item's metadata."""
    with logfire.span("api_update_knowledge_item") as span:
        span.set_attribute("endpoint", f"/api/knowledge-items/{source_id}")
        span.set_attribute("method", "PUT")
        span.set_attribute("source_id", source_id)
        
        try:
            logfire.info("Updating knowledge item", source_id=source_id, updates=updates)
            
            # Get Supabase client
            supabase_client = get_supabase_client()
            
            # Prepare the update data
            update_data = {}
            
            # Handle title updates
            if 'title' in updates:
                update_data['title'] = updates['title']
            
            # Handle description updates (stored in metadata)
            if 'description' in updates:
                # Get current metadata first
                current_response = supabase_client.table("sources").select("metadata").eq("source_id", source_id).execute()
                
                if current_response.data:
                    current_metadata = current_response.data[0].get('metadata', {})
                    current_metadata['description'] = updates['description']
                    update_data['metadata'] = current_metadata
                else:
                    # If no current metadata, create new
                    update_data['metadata'] = {'description': updates['description']}
            
            # Handle other metadata updates
            metadata_fields = ['knowledge_type', 'tags', 'status', 'update_frequency']
            for field in metadata_fields:
                if field in updates:
                    if 'metadata' not in update_data:
                        # Get current metadata first
                        current_response = supabase_client.table("sources").select("metadata").eq("source_id", source_id).execute()
                        current_metadata = current_response.data[0].get('metadata', {}) if current_response.data else {}
                        update_data['metadata'] = current_metadata
                    
                    update_data['metadata'][field] = updates[field]
            
            # Perform the update
            result = supabase_client.table("sources").update(update_data).eq("source_id", source_id).execute()
            
            if result.data:
                logfire.info("Knowledge item updated successfully", source_id=source_id)
                span.set_attribute("success", True)
                
                return {
                    'success': True,
                    'message': f'Successfully updated knowledge item {source_id}',
                    'source_id': source_id
                }
            else:
                logfire.error("Knowledge item not found", source_id=source_id)
                span.set_attribute("success", False)
                raise HTTPException(status_code=404, detail={'error': f'Knowledge item {source_id} not found'})
                
        except Exception as e:
            logfire.error("Failed to update knowledge item", error=str(e), source_id=source_id)
            span.set_attribute("error", str(e))
            raise HTTPException(status_code=500, detail={'error': str(e)})

@router.delete("/knowledge-items/{source_id}")
async def delete_knowledge_item(source_id: str):
    """Delete a knowledge item from the database."""
    with logfire.span("api_delete_knowledge_item") as span:
        span.set_attribute("endpoint", f"/api/knowledge-items/{source_id}")
        span.set_attribute("method", "DELETE")
        span.set_attribute("source_id", source_id)
        
        try:
            print(f"DEBUG: Starting delete_knowledge_item for source_id: {source_id}")
            logfire.info("Deleting knowledge item", source_id=source_id)
            
            # Get crawling context
            print(f"DEBUG: Getting crawling context...")
            crawling_context = get_crawling_context()
            print(f"DEBUG: Got crawling context: {type(crawling_context)}")
            
            # Ensure crawling context is initialized once
            print(f"DEBUG: Checking if context is initialized...")
            if not crawling_context._initialized:
                print(f"DEBUG: Initializing context...")
                await crawling_context.initialize()
                print(f"DEBUG: Context initialized")
            else:
                print(f"DEBUG: Context already initialized")
            
            # Create context for the MCP function
            print(f"DEBUG: Creating context...")
            ctx = crawling_context.create_context()
            print(f"DEBUG: Created context: {type(ctx)}")
            
            # Call the actual function from rag_module
            print(f"DEBUG: Importing delete_source function...")
            from src.modules.rag_module import delete_source_standalone as delete_source
            print(f"DEBUG: Successfully imported delete_source")
            
            print(f"DEBUG: Calling delete_source function...")
            result = await delete_source(ctx, source_id)
            print(f"DEBUG: delete_source returned: {result}")
            
            # Parse JSON string response if needed
            if isinstance(result, str):
                result = json.loads(result)
            
            if result.get('success'):
                logfire.info("Knowledge item deleted successfully", source_id=source_id)
                span.set_attribute("success", True)
                
                return {
                    'success': True,
                    'message': f'Successfully deleted knowledge item {source_id}'
                }
            else:
                logfire.error("Knowledge item deletion failed", source_id=source_id, error=result.get('error'))
                span.set_attribute("success", False)
                span.set_attribute("error", result.get('error'))
                raise HTTPException(status_code=500, detail={'error': result.get('error', 'Deletion failed')})
                
        except Exception as e:
            print(f"ERROR: Exception in delete_knowledge_item: {e}")
            print(f"ERROR: Exception type: {type(e)}")
            import traceback
            print(f"ERROR: Traceback: {traceback.format_exc()}")
            logfire.error("Failed to delete knowledge item", error=str(e), source_id=source_id)
            span.set_attribute("error", str(e))
            raise HTTPException(status_code=500, detail={'error': str(e)})

@router.post("/knowledge-items/crawl")
async def crawl_knowledge_item(request: KnowledgeItemRequest):
    """Crawl a URL and add it to the knowledge base with progress tracking."""
    with logfire.span("api_crawl_knowledge_item") as span:
        span.set_attribute("endpoint", "/api/knowledge-items/crawl")
        span.set_attribute("method", "POST")
        span.set_attribute("url", str(request.url))
        span.set_attribute("knowledge_type", request.knowledge_type)
        span.set_attribute("tags_count", len(request.tags))
        
        try:
            logfire.info("Starting knowledge item crawl", url=str(request.url), knowledge_type=request.knowledge_type, tags=request.tags)
            
            # Generate unique progress ID
            progress_id = str(uuid.uuid4())
            span.set_attribute("progress_id", progress_id)
            
            # Start progress tracking with initial state
            progress_manager.start_crawl(progress_id, {
                'progressId': progress_id,
                'currentUrl': str(request.url),
                'totalPages': 0,
                'processedPages': 0,
                'percentage': 0,
                'status': 'starting',
                'logs': [f'Starting crawl of {request.url}'],
                'eta': 'Calculating...'
            })
            
            # Start background task IMMEDIATELY (like the old API)
            asyncio.create_task(_perform_crawl_with_progress(progress_id, request))
            
            logfire.info("Crawl started successfully", progress_id=progress_id, url=str(request.url))
            span.set_attribute("success", True)
            
            response_data = {
                "success": True,
                "progressId": progress_id,
                "message": "Crawling started",
                "estimatedDuration": "3-5 minutes"
            }
            return response_data
            
        except Exception as e:
            logfire.error("Failed to start crawl", error=str(e), url=str(request.url))
            span.set_attribute("error", str(e))
            raise HTTPException(status_code=500, detail=str(e))

async def _perform_crawl_with_progress(progress_id: str, request: KnowledgeItemRequest):
    """Perform the actual crawl operation with progress tracking."""
    try:
        print(f"🚀 CRAWL: Starting crawl function with progress_id: {progress_id}")
        print(f"🚀 CRAWL: URL: {request.url}")
        
        # Wait for WebSocket connection to be established
        print(f"🚀 CRAWL: Waiting for WebSocket connection for progress_id: {progress_id}")
        connection_established = await progress_manager.wait_for_websocket_connection(progress_id, timeout=10.0)
        
        if not connection_established:
            print(f"⚠️ CRAWL: WebSocket connection not established in time for {progress_id}")
            # Continue anyway - progress will be stored and sent when connection is made
        else:
            print(f"✅ CRAWL: WebSocket connection verified for progress_id: {progress_id}")
        
        # Create a progress callback that will be called by the crawling function
        async def progress_callback(status: str, percentage: int, message: str, **kwargs):
            """Callback function to receive real-time progress updates from crawling."""
            print(f"🚀 CRAWL: Progress callback called with progress_id: {progress_id}")
            await progress_manager.update_progress(progress_id, {
                'status': status,
                'percentage': percentage,
                'currentUrl': kwargs.get('currentUrl', str(request.url)),
                'totalPages': kwargs.get('totalPages', 0),
                'processedPages': kwargs.get('processedPages', 0),
                'log': message,
                **kwargs
            })
            print(f"DEBUG: Progress callback - {status}: {percentage}% - {message}")
        
        # Initial progress update
        await progress_callback('starting', 0, f'Starting crawl of {request.url}')
        
        # Get crawling context
        crawling_context = get_crawling_context()
        
        # Ensure crawling context is initialized
        if not crawling_context._initialized:
            await crawling_context.initialize()
        
        # Create context for the MCP function
        ctx = crawling_context.create_context()
        
        # Store metadata in context for the crawling functions to access
        ctx.knowledge_metadata = {
            'knowledge_type': request.knowledge_type,
            'tags': request.tags,
            'update_frequency': request.update_frequency
        }
        
        # IMPORTANT: Add progress callback to context so MCP function can use it
        ctx.progress_callback = progress_callback
        
        print(f"🚀 CRAWL: About to call smart_crawl_url_direct with progress_id: {progress_id}")
        
        # Call the actual crawling function with progress callback support
        from src.modules.rag_module import smart_crawl_url_direct
        result = await smart_crawl_url_direct(
            ctx=ctx,
            url=str(request.url),
            max_depth=request.max_depth,
            max_concurrent=5,
            chunk_size=5000
        )
        
        print(f"🚀 CRAWL: smart_crawl_url_direct completed for progress_id: {progress_id}")
        
        # Parse JSON string response if needed
        if isinstance(result, str):
            result = json.loads(result)
        
        # Final completion update - the MCP function handles all phases including embeddings
        if result.get('success'):
            completion_data = {
                'chunksStored': result.get('chunks_stored', 0),
                'wordCount': result.get('total_word_count', 0),
                'log': 'All processing completed successfully (crawling, embeddings, and storage)'
            }
            print(f"🚀 CRAWL: All phases complete for progress_id: {progress_id}")
            await progress_manager.complete_crawl(progress_id, completion_data)
        else:
            print(f"🚀 CRAWL: Error in crawl for progress_id: {progress_id}")
            await progress_manager.error_crawl(progress_id, result.get('error', 'Unknown error'))
        
        # Broadcast final update to general WebSocket clients
        await manager.broadcast({
            "type": "crawl_completed",
            "data": {
                "url": str(request.url),
                "success": result.get('success', False),
                "message": f'Crawling completed for {request.url}',
                "progressId": progress_id
            }
        })
        
    except Exception as e:
        error_message = f'Crawling failed: {str(e)}'
        print(f"🚀 CRAWL: Exception in crawl for progress_id: {progress_id} - {error_message}")
        await progress_manager.error_crawl(progress_id, error_message)
        print(f"Crawl error for {progress_id}: {e}")

@router.post("/documents/upload")
async def upload_document(
    file: UploadFile = File(...),
    tags: Optional[str] = Form(None),
    knowledge_type: Optional[str] = Form("technical")
):
    """Upload and process a document with progress tracking."""
    with logfire.span("api_upload_document") as span:
        span.set_attribute("endpoint", "/api/documents/upload")
        span.set_attribute("method", "POST")
        span.set_attribute("filename", file.filename)
        span.set_attribute("content_type", file.content_type)
        span.set_attribute("knowledge_type", knowledge_type)
        
        try:
            logfire.info("Starting document upload", filename=file.filename, content_type=file.content_type, knowledge_type=knowledge_type)
            
            # Generate unique progress ID
            progress_id = str(uuid.uuid4())
            span.set_attribute("progress_id", progress_id)
            
            # Parse tags
            tag_list = json.loads(tags) if tags else []
            span.set_attribute("tags_count", len(tag_list))
            
            # Read file content immediately to avoid closed file issues
            file_content = await file.read()
            file_metadata = {
                'filename': file.filename,
                'content_type': file.content_type,
                'size': len(file_content)
            }
            
            # Start progress tracking
            progress_manager.start_crawl(progress_id, {
                'progressId': progress_id,
                'status': 'starting',
                'percentage': 0,
                'currentUrl': f"file://{file.filename}",
                'logs': [f'Starting upload of {file.filename}'],
                'uploadType': 'document',
                'fileName': file.filename,
                'fileType': file.content_type
            })
            
            # Start background task for processing with file content and metadata
            asyncio.create_task(_perform_upload_with_progress(progress_id, file_content, file_metadata, tag_list, knowledge_type))
            
            logfire.info("Document upload started successfully", progress_id=progress_id, filename=file.filename)
            span.set_attribute("success", True)
            
            return {
                "success": True,
                "progressId": progress_id,
                "message": "Document upload started",
                "filename": file.filename
            }
            
        except Exception as e:
            logfire.error("Failed to start document upload", error=str(e), filename=file.filename, error_type=type(e).__name__)
            span.set_attribute("error", str(e))
            span.set_attribute("error_type", type(e).__name__)
            raise HTTPException(status_code=500, detail={'error': str(e)})

async def _perform_upload_with_progress(progress_id: str, file_content: bytes, file_metadata: dict, tag_list: List[str], knowledge_type: str):
    """Perform document upload with progress tracking using callback pattern like web crawl."""
    try:
        filename = file_metadata['filename']
        content_type = file_metadata['content_type']
        file_size = file_metadata['size']
        
        # Wait for WebSocket connection to be established
        print(f"🚀 UPLOAD: Waiting for WebSocket connection for progress_id: {progress_id}")
        connection_established = await progress_manager.wait_for_websocket_connection(progress_id, timeout=10.0)
        
        if not connection_established:
            print(f"⚠️ UPLOAD: WebSocket connection not established in time for {progress_id}")
            # Continue anyway - progress will be stored and sent when connection is made
        else:
            print(f"✅ UPLOAD: WebSocket connection verified for progress_id: {progress_id}")
        
        # Create a progress callback function exactly like web crawl
        async def progress_callback(status: str, percentage: int, message: str, **kwargs):
            """Callback function to send progress updates - matches web crawl pattern."""
            print(f"🚀 UPLOAD: Progress callback called with progress_id: {progress_id}")
            await progress_manager.update_progress(progress_id, {
                'status': status,
                'percentage': percentage,
                'currentUrl': f"file://{filename}",
                'log': message,
                **kwargs
            })
            print(f"DEBUG: Upload progress callback - {status}: {percentage}% - {message}")
        
        # Step 1: Reading file (already done, but report it)
        await progress_callback('reading', 10, f'File content loaded ({file_size} bytes)...')
        
        # Step 2: Extract text - DO THE WORK THEN REPORT
        extracted_text = None
        try:
            extracted_text = extract_text_from_document(file_content, filename, content_type)
            logfire.info("Document text extracted successfully", 
                       filename=filename, 
                       extracted_length=len(extracted_text),
                       content_type=content_type)
            # ONLY report progress AFTER the work is done
            await progress_callback('extracting', 25, f'Text extracted successfully ({len(extracted_text)} characters)')
        except Exception as e:
            await progress_manager.error_crawl(progress_id, f"Failed to extract text: {str(e)}")
            return
        
        # Step 3: Chunk the text - DO THE WORK THEN REPORT
        storage_service = DocumentStorageService(get_supabase_client())
        chunks = storage_service.smart_chunk_markdown(extracted_text, chunk_size=5000)
        logfire.info("Document chunked", filename=filename, chunk_count=len(chunks))
        # ONLY report progress AFTER the work is done
        await progress_callback('chunking', 40, f'Document broken into {len(chunks)} chunks')
        
        # Step 4: Create source - DO THE WORK THEN REPORT
        supabase_client = get_supabase_client()
        source_id = f"upload_{int(datetime.now().timestamp())}"
        total_word_count = sum(len(chunk.split()) for chunk in chunks)
        
        # ONLY report progress AFTER source ID is generated
        await progress_callback('creating_source', 55, f'Creating source entry for {len(chunks)} chunks...')
        
        # Step 5: Generate summary - DO THE WORK THEN REPORT
        source_summary = extract_source_summary(source_id, extracted_text[:5000])
        
        source_data = {
            "source_id": source_id,
            "title": filename,
            "summary": source_summary,
            "metadata": {
                "knowledge_type": knowledge_type,
                "tags": tag_list,
                "source_type": "file",
                "file_name": filename,
                "file_type": content_type,
                "file_size": file_size,
                "chunk_count": len(chunks)
            },
            "total_word_count": total_word_count,
            "update_frequency": 0  # Files don't auto-update, set to never
        }
        
        sources_response = supabase_client.table("sources").insert(source_data).execute()
        logfire.info("Source created successfully", source_id=source_id)
        # ONLY report progress AFTER source is created
        await progress_callback('summarizing', 65, f'AI summary generated and source created')
        
        # Step 6: Store chunks with REAL-TIME progress reporting during storage
        # Progress will be reported by add_documents_to_supabase through scaled callback
        
        documents = []
        for i, chunk in enumerate(chunks):
            section_info = storage_service.extract_section_info(chunk)
            
            document_data = {
                "url": f"file://{filename}",
                "content": chunk,
                "source_id": source_id,
                "chunk_number": i,
                "metadata": {
                    "title": filename,
                    "knowledge_type": knowledge_type,
                    "tags": tag_list,
                    "headers": section_info["headers"],
                    "char_count": section_info["char_count"],
                    "word_count": section_info["word_count"],
                    "file_name": filename,
                    "file_type": content_type,
                    "chunk_index": i,
                    "source_type": "file"
                }
            }
            documents.append(document_data)
        
        # Store documents with INCREMENTAL progress reporting during the process
        try:
            urls = [doc['url'] for doc in documents]
            chunk_numbers = [doc['chunk_number'] for doc in documents]
            contents = [doc['content'] for doc in documents]
            metadatas = [doc['metadata'] for doc in documents]
            url_to_full_document = {f"file://{filename}": extracted_text}
            
            # CRITICAL FIX: Use the working function with progress reporting during storage
            # Instead of calling add_documents_to_supabase silently, we'll call it with progress
            from src.server.utils import add_documents_to_supabase
            
            # Create a scaled progress callback that maps 0-100% to 65-100%
            async def scaled_progress_callback(status: str, percentage: int, message: str, **kwargs):
                # Scale the percentage from document storage (0-100) to overall progress (65-100)
                scaled_percentage = 65 + int((percentage / 100.0) * 35)
                # Keep status consistent as 'summarizing' to avoid frontend progress bar reset
                await progress_callback('summarizing', scaled_percentage, message, **kwargs)
            
            # Call add_documents_to_supabase with scaled progress callback
            await add_documents_to_supabase(
                supabase_client, 
                urls, 
                chunk_numbers, 
                contents, 
                metadatas, 
                url_to_full_document,
                batch_size=20,
                progress_callback=scaled_progress_callback
            )
            
            # Progress already reported to 100% by scaled callback
            
            logfire.info("Document chunks inserted successfully", 
                       source_id=source_id, 
                       chunk_count=len(chunks),
                       total_word_count=total_word_count)
            
        except Exception as supabase_error:
            await progress_manager.error_crawl(progress_id, f"Failed to store chunks: {str(supabase_error)}")
            return
        
        # FINAL COMPLETION - only called after ALL work is done
        await progress_manager.complete_crawl(progress_id, {
            'chunksStored': len(chunks),
            'wordCount': total_word_count,
            'sourceId': source_id,
            'log': f'Document upload completed successfully! {len(chunks)} chunks stored.'
        })
        
        # Broadcast to general WebSocket clients
        await manager.broadcast({
            "type": "upload_completed",
            "data": {
                "filename": filename,
                "success": True,
                "message": f'Document upload completed for {filename}',
                "progressId": progress_id,
                "sourceId": source_id
            }
        })
        
        logfire.info("Document uploaded successfully", source_id=source_id, filename=filename, file_size=file_size)
        
    except Exception as e:
        logfire.error("Failed to upload document", error=str(e), filename=file_metadata.get('filename', 'unknown'), error_type=type(e).__name__)
        await progress_manager.error_crawl(progress_id, f"Upload failed: {str(e)}")
        
        # Broadcast error to general WebSocket clients
        await manager.broadcast({
            "type": "upload_error",
            "data": {
                "filename": file_metadata.get('filename', 'unknown'),
                "success": False,
                "error": str(e),
                "progressId": progress_id
            }
        })

@router.post("/rag/query")
async def perform_rag_query(request: RagQueryRequest):
    """Perform a RAG query on the knowledge base."""
    with logfire.span("api_rag_query") as span:
        span.set_attribute("endpoint", "/api/rag/query")
        span.set_attribute("method", "POST")
        span.set_attribute("query_length", len(request.query))
        span.set_attribute("source", request.source)
        span.set_attribute("match_count", request.match_count)
        
        try:
            logfire.info("RAG query started", 
                           query=request.query[:100] + "..." if len(request.query) > 100 else request.query,
                           source=request.source,
                           match_count=request.match_count)
            
            # Use simple text matching for FastAPI direct query
            supabase = get_supabase_client()
            
            with logfire.span("database_query"):
                query_builder = supabase.table("crawled_pages").select("*")
                
                # Apply source filter if provided
                if request.source:
                    logfire.debug("Applying source filter", source=request.source)
                    query_builder = query_builder.ilike("source", f"%{request.source}%")
                
                # Apply text search
                query_builder = query_builder.ilike("content", f"%{request.query}%")
                query_builder = query_builder.limit(request.match_count)
                
                response = query_builder.execute()
            
            results = response.data if response.data else []
            
            # Format results consistently
            with logfire.span("format_results"):
                formatted_results = []
                for result in results:
                    formatted_results.append({
                        "id": result.get("id"),
                        "content": result.get("content", "")[:1000],  # Limit content length
                        "metadata": {
                            "source": result.get("source"),
                            "url": result.get("url"),
                            "title": result.get("title"),
                            "chunk_index": result.get("chunk_index", 0)
                        },
                        "score": 0.8  # Mock similarity score for text matching
                    })
            
            span.set_attribute("results_count", len(formatted_results))
            logfire.info("RAG query completed successfully", 
                           results_count=len(formatted_results),
                           execution_path="fastapi_direct")
            
            return {
                "success": True,
                "results": formatted_results,
                "query": request.query,
                "source": request.source,
                "match_count": request.match_count,
                "execution_path": "fastapi_direct"
            }
            
        except Exception as e:
            logfire.error("RAG query failed", error=str(e), query=request.query[:50], source=request.source)
            span.set_attribute("error", str(e))
            raise HTTPException(status_code=500, detail=f"RAG query failed: {str(e)}")

@router.get("/rag/sources")
async def get_available_sources():
    """Get all available sources for RAG queries."""
    with logfire.span("api_get_available_sources") as span:
        span.set_attribute("endpoint", "/api/rag/sources")
        span.set_attribute("method", "GET")
        
        try:
            logfire.info("Getting available RAG sources")
            result = await get_available_sources_direct()
            
            # Parse result if it's a string
            if isinstance(result, str):
                result = json.loads(result)
            
            sources_count = len(result.get('sources', []))
            logfire.info("Available RAG sources retrieved", sources_count=sources_count)
            span.set_attribute("sources_count", sources_count)
            span.set_attribute("success", result.get('success', False))
            
            return result
            
        except Exception as e:
            logfire.error("Failed to get available sources", error=str(e))
            span.set_attribute("error", str(e))
            raise HTTPException(status_code=500, detail={'error': str(e)})

# WebSocket Endpoints
@router.websocket("/knowledge-items/stream")
async def websocket_knowledge_items(websocket: WebSocket):
    """WebSocket endpoint for real-time knowledge items updates."""
    await manager.connect(websocket)
    try:
        # Ensure crawling context is initialized once  
        crawling_context = get_crawling_context()
        if not crawling_context._initialized:
            await crawling_context.initialize()
        
        # Send initial data
        sources_result = await get_available_sources_direct()
        
        if isinstance(sources_result, str):
            sources_data = json.loads(sources_result)
        else:
            sources_data = sources_result
        
        # Transform data for frontend (use same mapping as REST API)
        items = []
        for source in sources_data.get('sources', []):
            # Use metadata from sources table
            source_metadata = source.get('metadata', {})
            
            # Get first page URL if available
            supabase_client = crawling_context.supabase_client or get_supabase_client()
            pages_response = supabase_client.from_('crawled_pages')\
                .select('url')\
                .eq('source_id', source['source_id'])\
                .limit(1)\
                .execute()
            
            first_page = pages_response.data[0] if pages_response.data else {}
            
            # Determine source type - if metadata has source_type='file', use it; otherwise check URL pattern
            stored_source_type = source_metadata.get('source_type')
            if stored_source_type:
                source_type = stored_source_type
            else:
                # Legacy fallback - check URL pattern
                first_page_url = first_page.get('url', f"source://{source['source_id']}")
                source_type = 'file' if first_page_url.startswith('file://') else 'url'
            
            item = {
                'id': source['source_id'],
                'title': source.get('title', source.get('summary', 'Untitled')),
                'url': first_page.get('url', f"source://{source['source_id']}"),
                'source_id': source['source_id'],
                'metadata': {
                    'knowledge_type': source_metadata.get('knowledge_type', 'technical'),
                    'tags': source_metadata.get('tags', []),
                    'source_type': source_type,
                    'status': 'active',
                    'description': source_metadata.get('description', source.get('summary', '')),
                    'chunks_count': source.get('total_words', 0),
                    'word_count': source.get('total_words', 0),
                    'last_scraped': source.get('updated_at'),
                    'file_name': source_metadata.get('file_name'),
                    'file_type': source_metadata.get('file_type'),
                    'update_frequency': source.get('update_frequency', 7),  # Include frequency from sources table
                    **source_metadata
                },
                'created_at': source.get('created_at'),
                'updated_at': source.get('updated_at')
            }
            items.append(item)
        
        await websocket.send_json({
            "type": "knowledge_items_update",
            "data": {
                "items": items,
                "total": len(items),
                "page": 1,
                "per_page": 20,
                "pages": 1
            }
        })
        
        # Keep connection alive and listen for updates
        while True:
            await asyncio.sleep(5)  # Check for updates every 5 seconds
            # Send heartbeat
            await websocket.send_json({"type": "heartbeat"})
            
    except WebSocketDisconnect:
        manager.disconnect(websocket)
    except Exception as e:
        print(f"WebSocket error: {e}")
        manager.disconnect(websocket)

@router.websocket("/crawl-progress/{progress_id}")
async def websocket_crawl_progress(websocket: WebSocket, progress_id: str):
    """WebSocket endpoint for tracking specific crawl progress."""
    if USE_SOCKETIO_PROGRESS:
        # Socket.IO is enabled, reject WebSocket connection with informative message
        await websocket.accept()
        await websocket.send_json({
            "type": "error",
            "message": "WebSocket endpoint deprecated. Please use Socket.IO connection on /crawl namespace."
        })
        await websocket.close()
        return
    
    try:
        print(f"DEBUG: WebSocket connecting for progress_id: {progress_id}")
        print(f"DEBUG: Active crawls: {list(progress_manager.active_crawls.keys())}")
        print(f"DEBUG: WebSocket headers: {websocket.headers}")
        
        # Add WebSocket to progress manager (handles accept internally)
        await progress_manager.add_websocket(progress_id, websocket)
        
        print(f"DEBUG: WebSocket registered for progress_id: {progress_id}")
        
        # Keep connection alive with improved timeout handling
        heartbeat_interval = 30.0  # Send heartbeat every 30 seconds
        client_timeout = 60.0      # Wait up to 60 seconds for client messages
        
        while True:
            try:
                # Wait for messages from client (like ping) with timeout
                message = await asyncio.wait_for(
                    websocket.receive_text(), 
                    timeout=heartbeat_interval
                )
                
                print(f"DEBUG: Received client message: {message}")
                
                # Respond to ping messages
                if message == "ping":
                    await websocket.send_json({"type": "pong"})
                    print(f"DEBUG: Sent pong response to client")
                    
            except asyncio.TimeoutError:
                # Send heartbeat every heartbeat_interval seconds
                try:
                    await websocket.send_json({"type": "heartbeat"})
                    print(f"DEBUG: Sent heartbeat for progress_id: {progress_id}")
                except Exception as e:
                    print(f"DEBUG: Failed to send heartbeat, connection likely dead: {e}")
                    break
                    
            except WebSocketDisconnect:
                print(f"DEBUG: Client disconnected normally for progress_id: {progress_id}")
                break
                
            except Exception as e:
                print(f"DEBUG: Unexpected error in WebSocket loop: {e}")
                break
                
    except WebSocketDisconnect:
        print(f"DEBUG: WebSocket disconnected for progress_id: {progress_id}")
    except Exception as e:
        print(f"DEBUG: WebSocket error for progress_id {progress_id}: {e}")
        import traceback
        traceback.print_exc()
    finally:
        print(f"DEBUG: Cleaning up WebSocket for progress_id: {progress_id}")
        progress_manager.remove_websocket(progress_id, websocket)

@router.get("/database/metrics")
async def get_database_metrics():
    """Get database metrics and statistics."""
    with logfire.span("api_get_database_metrics") as span:
        span.set_attribute("endpoint", "/api/database/metrics")
        span.set_attribute("method", "GET")
        
        try:
            logfire.info("Getting database metrics")
            supabase_client = get_supabase_client()
            
            # Get counts from various tables
            sources_count = supabase_client.table("sources").select("*", count="exact").execute()
            pages_count = supabase_client.table("crawled_pages").select("*", count="exact").execute()
            
            sources_count_val = sources_count.count if sources_count.count else 0
            pages_count_val = pages_count.count if pages_count.count else 0
            
            logfire.info("Database metrics retrieved", 
                        sources_count=sources_count_val, 
                        pages_count=pages_count_val)
            span.set_attribute("sources_count", sources_count_val)
            span.set_attribute("pages_count", pages_count_val)
            
            return {
                "sources_count": sources_count_val,
                "pages_count": pages_count_val,
                "timestamp": datetime.now().isoformat()
            }
            
        except Exception as e:
            logfire.error("Failed to get database metrics", error=str(e))
            span.set_attribute("error", str(e))
            raise HTTPException(status_code=500, detail={'error': str(e)})

@router.get("/health")
async def knowledge_health():
    """Knowledge API health check."""
    with logfire.span("api_knowledge_health") as span:
        span.set_attribute("endpoint", "/api/health")
        span.set_attribute("method", "GET")
        
        # Removed health check logging to reduce console noise
        result = {
            "status": "healthy",
            "service": "knowledge-api",
            "timestamp": datetime.now().isoformat()
        }
        span.set_attribute("status", "healthy")
        
        return result 

def extract_text_from_document(file_content: bytes, filename: str, content_type: str) -> str:
    """
    Extract text from various document formats.
    
    Args:
        file_content: Raw file bytes
        filename: Name of the file
        content_type: MIME type of the file
        
    Returns:
        Extracted text content
        
    Raises:
        ValueError: If the file format is not supported
        Exception: If extraction fails
    """
    try:
        # PDF files
        if content_type == 'application/pdf' or filename.lower().endswith('.pdf'):
            return extract_text_from_pdf(file_content)
        
        # Word documents
        elif (content_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document', 'application/msword'] 
              or filename.lower().endswith(('.docx', '.doc'))):
            return extract_text_from_docx(file_content)
        
        # Text files (markdown, txt, etc.)
        elif (content_type.startswith('text/') 
              or filename.lower().endswith(('.txt', '.md', '.markdown', '.rst'))):
            return file_content.decode('utf-8', errors='ignore')
        
        else:
            raise ValueError(f"Unsupported file format: {content_type} ({filename})")
            
    except Exception as e:
        logfire.error("Document text extraction failed", 
                     filename=filename, 
                     content_type=content_type, 
                     error=str(e))
        raise Exception(f"Failed to extract text from {filename}: {str(e)}")

def extract_text_from_pdf(file_content: bytes) -> str:
    """
    Extract text from PDF using both PyPDF2 and pdfplumber for best results.
    
    Args:
        file_content: Raw PDF bytes
        
    Returns:
        Extracted text content
    """
    if not PDFPLUMBER_AVAILABLE and not PYPDF2_AVAILABLE:
        raise Exception("No PDF processing libraries available. Please install pdfplumber and PyPDF2.")
    
    text_content = []
    
    # First try with pdfplumber (better for complex layouts)
    if PDFPLUMBER_AVAILABLE:
        try:
            with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text_content.append(f"--- Page {page_num + 1} ---\n{page_text}")
                    except Exception as e:
                        logfire.warning(f"pdfplumber failed on page {page_num + 1}: {e}")
                        continue
            
            # If pdfplumber got good results, use them
            if text_content and len('\n'.join(text_content).strip()) > 100:
                return '\n\n'.join(text_content)
            
        except Exception as e:
            logfire.warning(f"pdfplumber extraction failed: {e}, trying PyPDF2")
    
    # Fallback to PyPDF2
    if PYPDF2_AVAILABLE:
        try:
            text_content = []
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
            
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text_content.append(f"--- Page {page_num + 1} ---\n{page_text}")
                except Exception as e:
                    logfire.warning(f"PyPDF2 failed on page {page_num + 1}: {e}")
                    continue
            
            if text_content:
                return '\n\n'.join(text_content)
            else:
                raise Exception("No text could be extracted from PDF")
                
        except Exception as e:
            raise Exception(f"PyPDF2 failed to extract text: {str(e)}")
    
    # If we get here, no libraries worked
    raise Exception("Failed to extract text from PDF - no working PDF libraries available")

def extract_text_from_docx(file_content: bytes) -> str:
    """
    Extract text from Word documents (.docx).
    
    Args:
        file_content: Raw DOCX bytes
        
    Returns:
        Extracted text content
    """
    if not DOCX_AVAILABLE:
        raise Exception("python-docx library not available. Please install python-docx.")
    
    try:
        doc = DocxDocument(io.BytesIO(file_content))
        text_content = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_content.append(' | '.join(row_text))
        
        if not text_content:
            raise Exception("No text content found in document")
            
        return '\n\n'.join(text_content)
        
    except Exception as e:
        raise Exception(f"Failed to extract text from Word document: {str(e)}")

# Socket.IO Event Handlers for Crawl Progress
from ..socketio_app import get_socketio_instance, NAMESPACE_CRAWL

sio = get_socketio_instance()

@sio.on('connect', namespace=NAMESPACE_CRAWL)
async def on_crawl_connect(sid, environ):
    """Handle Socket.IO connection for crawl progress."""
    print(f"🔌 Crawl client connected: {sid}")
    await sio.emit('connected', {'message': 'Connected to crawl progress'}, to=sid, namespace=NAMESPACE_CRAWL)

@sio.on('disconnect', namespace=NAMESPACE_CRAWL)
async def on_crawl_disconnect(sid):
    """Handle Socket.IO disconnection."""
    print(f"🔌 Crawl client disconnected: {sid}")

@sio.on('subscribe', namespace=NAMESPACE_CRAWL)
async def on_subscribe_progress(sid, data):
    """Subscribe to crawl progress updates."""
    progress_id = data.get('progress_id')
    if not progress_id:
        await sio.emit('error', {'message': 'progress_id required'}, to=sid, namespace=NAMESPACE_CRAWL)
        return
    
    # Join the room for this progress ID
    await sio.enter_room(sid, progress_id, namespace=NAMESPACE_CRAWL)
    print(f"✅ Client {sid} subscribed to progress {progress_id}")

@sio.on('unsubscribe', namespace=NAMESPACE_CRAWL)
async def on_unsubscribe_progress(sid, data):
    """Unsubscribe from crawl progress updates."""
    progress_id = data.get('progress_id')
    if progress_id:
        await sio.leave_room(sid, progress_id, namespace=NAMESPACE_CRAWL)
        print(f"🛑 Client {sid} unsubscribed from progress {progress_id}") 
